from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Literal

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader

SourceType = Literal["pdf", "docx", "html", "markdown", "text", "log"]


@dataclass
class ParsedSection:
    heading: str
    text: str
    page: int | None = None


@dataclass
class ParsedDocument:
    source_type: SourceType
    filename: str
    sections: list[ParsedSection]
    raw_text: str


def detect_source_type(filename: str, content_type: str | None = None) -> SourceType:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf" or (content_type and "pdf" in content_type):
        return "pdf"
    if ext == ".docx" or (content_type and "wordprocessingml" in (content_type or "")):
        return "docx"
    if ext in {".html", ".htm"} or (content_type and "html" in content_type):
        return "html"
    if ext in {".md", ".markdown"}:
        return "markdown"
    if ext == ".log":
        return "log"
    return "text"


def _split_by_headings(lines: list[str], is_heading) -> list[ParsedSection]:
    sections: list[ParsedSection] = []
    current_heading = "Introduction"
    buffer: list[str] = []
    for line in lines:
        if is_heading(line):
            if buffer:
                sections.append(ParsedSection(heading=current_heading, text="\n".join(buffer).strip()))
                buffer = []
            current_heading = line.strip().lstrip("#").strip() or current_heading
        else:
            buffer.append(line)
    if buffer:
        sections.append(ParsedSection(heading=current_heading, text="\n".join(buffer).strip()))
    return [s for s in sections if s.text]


def parse_pdf(data: bytes, filename: str) -> ParsedDocument:
    reader = PdfReader(io.BytesIO(data))
    all_text_parts: list[str] = []
    sections: list[ParsedSection] = []
    for page_num, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        all_text_parts.append(text)
        # Very lightweight section detection: uppercase or numbered lines
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        current = "Page " + str(page_num)
        buf: list[str] = []

        def _is_head(l: str) -> bool:
            if len(l) > 80:
                return False
            if re.match(r"^\d+(\.\d+)*\s+[A-Z]", l):
                return True
            if l.isupper() and len(l) > 3:
                return True
            return False

        for l in lines:
            if _is_head(l):
                if buf:
                    sections.append(ParsedSection(heading=current, text="\n".join(buf).strip(), page=page_num))
                    buf = []
                current = l
            else:
                buf.append(l)
        if buf:
            sections.append(ParsedSection(heading=current, text="\n".join(buf).strip(), page=page_num))
    raw = "\n\n".join(all_text_parts)
    return ParsedDocument(source_type="pdf", filename=filename, sections=sections, raw_text=raw)


def parse_docx(data: bytes, filename: str) -> ParsedDocument:
    doc = DocxDocument(io.BytesIO(data))
    lines: list[str] = []
    headings_style_set = {"Heading 1", "Heading 2", "Heading 3", "Title"}
    heading_flags: list[bool] = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        lines.append(txt)
        heading_flags.append(p.style.name in headings_style_set)
    sections: list[ParsedSection] = []
    buf: list[str] = []
    current = "Introduction"
    for line, is_head in zip(lines, heading_flags):
        if is_head:
            if buf:
                sections.append(ParsedSection(heading=current, text="\n".join(buf).strip()))
                buf = []
            current = line
        else:
            buf.append(line)
    if buf:
        sections.append(ParsedSection(heading=current, text="\n".join(buf).strip()))
    return ParsedDocument(
        source_type="docx", filename=filename, sections=sections, raw_text="\n".join(lines)
    )


def parse_html(data: bytes, filename: str) -> ParsedDocument:
    soup = BeautifulSoup(data, "html.parser")
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()
    sections: list[ParsedSection] = []
    current = soup.title.string.strip() if soup.title and soup.title.string else "Document"
    buf: list[str] = []
    for el in soup.body.descendants if soup.body else []:
        if getattr(el, "name", None) in {"h1", "h2", "h3", "h4"}:
            if buf:
                sections.append(ParsedSection(heading=current, text="\n".join(buf).strip()))
                buf = []
            current = el.get_text(strip=True) or current
        elif hasattr(el, "get_text"):
            continue
        else:
            text = str(el).strip()
            if text:
                buf.append(text)
    if not sections:
        text = soup.get_text(separator="\n")
        sections.append(ParsedSection(heading=current, text=re.sub(r"\n{3,}", "\n\n", text).strip()))
    raw = re.sub(r"\n{3,}", "\n\n", soup.get_text(separator="\n")).strip()
    return ParsedDocument(source_type="html", filename=filename, sections=sections, raw_text=raw)


def parse_markdown(data: bytes, filename: str) -> ParsedDocument:
    text = data.decode("utf-8", errors="replace")
    lines = text.splitlines()
    sections = _split_by_headings(lines, lambda l: l.strip().startswith("#"))
    return ParsedDocument(source_type="markdown", filename=filename, sections=sections, raw_text=text)


def parse_text(data: bytes, filename: str) -> ParsedDocument:
    text = data.decode("utf-8", errors="replace")
    lines = text.splitlines()
    sections = _split_by_headings(lines, lambda l: bool(re.match(r"^(?:#+\s|[A-Z][A-Z0-9 ]{3,}$)", l.strip())))
    if not sections:
        sections = [ParsedSection(heading="Document", text=text.strip())]
    return ParsedDocument(source_type="text", filename=filename, sections=sections, raw_text=text)


def parse(data: bytes, filename: str, content_type: str | None = None) -> ParsedDocument:
    stype = detect_source_type(filename, content_type)
    if stype == "pdf":
        return parse_pdf(data, filename)
    if stype == "docx":
        return parse_docx(data, filename)
    if stype == "html":
        return parse_html(data, filename)
    if stype == "markdown":
        return parse_markdown(data, filename)
    return parse_text(data, filename)
